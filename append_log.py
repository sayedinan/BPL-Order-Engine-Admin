"""
Append a subtask entry to SUBTASKS log.docx.

Usage:
    python append_log.py "<task-id>" "<subtask-id>" "<description>" "<files>" [status]

Examples:
    python append_log.py "1.1.1" "1.1.1" "Add react-router-dom@7 to package.json" "frontend/package.json"
    python append_log.py "1.4.1" "1.4.1" "Create AuthContext with v0.3 shape" "src/auth/AuthContext.tsx"

The entry is appended as a new row to the subtask log table at the end of the document.
If the document doesn't exist yet, it is created with a header table.
"""
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DOCX_PATH = Path("SUBTASKS log.docx")


def shade_cell(cell, fill_hex: str) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def ensure_document() -> Document:
    """Open the existing docx, or create one with a header table if it doesn't exist."""
    if DOCX_PATH.exists():
        return Document(str(DOCX_PATH))

    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("BPL Order Engine Admin — Subtask Log", level=0)

    # Intro paragraph
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Live subtask log. One row per subtask landed. Created "
        + datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        + "."
    )
    intro_run.italic = True
    intro_run.font.size = Pt(10)
    intro_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Header table
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Timestamp (UTC)", "Task", "Subtask", "Description", "Files", "Status"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(cell, "D9E1F2")

    return doc


def append_entry(
    task_id: str,
    subtask_id: str,
    description: str,
    files: str,
    status: str = "✅",
) -> None:
    import gc

    doc = ensure_document()

    # Find the last table in the document (assumed to be the log table).
    if not doc.tables:
        # Shouldn't happen because ensure_document created one, but defensive.
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Timestamp (UTC)", "Task", "Subtask", "Description", "Files", "Status"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
            shade_cell(cell, "D9E1F2")
    else:
        table = doc.tables[-1]

    row = table.add_row()
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
    cells = row.cells
    cells[0].text = timestamp
    cells[1].text = task_id
    cells[2].text = subtask_id
    cells[3].text = description
    cells[4].text = files
    cells[5].text = status

    # Color-code by status
    if status.lower().startswith("✅") or status.lower() == "done":
        shade_cell(cells[5], "D5F5E3")  # green
    elif status.lower().startswith("⚠") or status.lower() == "blocked":
        shade_cell(cells[5], "FCE4D6")  # orange
    elif status.lower().startswith("❌") or status.lower() == "failed":
        shade_cell(cells[5], "F8CBAD")  # red
    else:
        shade_cell(cells[5], "FFF2CC")  # yellow (in progress)

    # Force a GC so any lingering ZipFile handle on Windows is released
    # before the next save. python-docx keeps an internal Part object
    # that can hold the file open; this is a known sharp edge.
    doc_path = Path(DOCX_PATH)
    # Save to a temp path then move to avoid Windows file lock races
    # when called in quick succession.
    tmp_path = doc_path.with_suffix(".docx.tmp")
    doc.save(str(tmp_path))
    # Force-close the doc object to release internal handles
    del doc
    gc.collect()
    # Atomic replace.
    if doc_path.exists():
        doc_path.unlink()
    tmp_path.rename(doc_path)
    print(f"Appended: {subtask_id} — {description[:60]}…")


if __name__ == "__main__":
    if len(sys.argv) < 5:
        print(
            "Usage: python append_log.py <task-id> <subtask-id> <description> <files> [status]",
            file=sys.stderr,
        )
        sys.exit(1)

    task_id = sys.argv[1]
    subtask_id = sys.argv[2]
    description = sys.argv[3]
    files = sys.argv[4]
    status = sys.argv[5] if len(sys.argv) > 5 else "✅"

    append_entry(task_id, subtask_id, description, files, status)
