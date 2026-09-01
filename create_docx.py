"""
Create a DOCX file documenting all the subtasks performed in the
BPL-Order-Engine-Admin project (as decomposed in TASKS-decomposed.md).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, fill_hex):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_subtask_table(doc, subtasks):
    """Add a table for subtask list. Each row: [#, Subtask, Touches, Done when]"""
    table = doc.add_table(rows=1 + len(subtasks), cols=4)
    table.style = "Light Grid Accent 1"

    # Header
    headers = ["#", "Subtask", "Touches", "Done when"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(hdr_cells[i], "D9E1F2")

    for idx, (num, subtask, touches, done_when, is_secure) in enumerate(subtasks, start=1):
        row = table.rows[idx].cells
        row[0].text = num
        row[1].text = subtask
        row[2].text = touches
        row[3].text = done_when
        if is_secure:
            shade_cell(row[0], "FFE4E1")  # pinkish for security-relevant


def add_summary_table(doc, summary):
    table = doc.add_table(rows=1 + len(summary), cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Task", "Subtasks", "🔒 count"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(hdr_cells[i], "D9E1F2")

    for idx, (task, sub, sec) in enumerate(summary, start=1):
        row = table.rows[idx].cells
        row[0].text = task
        row[1].text = str(sub)
        row[2].text = str(sec)


# ============================================================
# Build the document
# ============================================================
doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---- Title ----
title = doc.add_heading("BPL-Order-Engine-Admin — Subtask Documentation", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documented work performed in the project (per TASKS-decomposed.md)")
run.italic = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date: September 1, 2026  •  Orchestrator: Sayed")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ---- Overview ----
add_heading(doc, "Overview", level=1)
add_paragraph(
    doc,
    "This document records the subtasks performed in the BPL-Order-Engine-Admin project. "
    "It is generated from TASKS-decomposed.md, which was produced via the "
    "task-decomposition skill. Every task in TASKS.md was broken down into a table of "
    "small, paste-able subtasks. The subtasks below describe what was done (or planned) "
    "under each parent task, with explicit 'Done when' gates and the files each subtask "
    "touched. Security-relevant seams (auth, RBAC, credentials, SSH) are marked with 🔒.",
)
add_paragraph(
    doc,
    "Legend: 🔒 = security-relevant seam (auth, RBAC, credentials, SSH). Extra scrutiny in review.",
    italic=True,
)

# ---- Task #11 ----
add_heading(doc, "Task #11 — SPEC.md USER audit log fix  📄", level=2)
add_paragraph(doc, "Parent task: TASKS.md #11.  Files touched: SPEC.md (one file, one commit).", italic=True)
add_subtask_table(doc, [
    ("11.1", "Edit §3.1 RBAC matrix `View audit logs` row: USER column from `🔒 assigned engines only` to `❌`",
     "SPEC.md (one row, line 324)", "`grep -n \"View audit logs\" SPEC.md` shows the USER cell is `❌`", False),
    ("11.2", "Edit §4.5 `GET /api/audit-logs` heading: from \"depends on role\" to \"`SYS_ADMIN`, `ADMIN` only\"",
     "SPEC.md (one heading, line 580)", "`grep -n \"GET /api/audit-logs\" SPEC.md` shows the new heading", False),
    ("11.3", "Edit §4.5 USER clause: replace the assignedEngines-filtered bullet with a 403 note",
     "SPEC.md (one bullet, line 583)", "`grep -c \"403\" SPEC.md` is `>= 2` (new note + envelope ref)", False),
    ("11.4", "Edit §5.4 Logs page: add a sentence that the Source dropdown hides \"System Audit Logs\" for USER",
     "SPEC.md (one paragraph, ~line 671)", "`grep -n \"Source dropdown\" SPEC.md` returns the new sentence", False),
    ("11.5", "Edit §5.1 routes table: keep `/logs` as \"all authenticated\" but add a note that USER sees a reduced view",
     "SPEC.md (one row in the routes table)", "`grep -n \"reduced view\" SPEC.md` returns the new note", False),
])

# ---- Task #12 ----
add_heading(doc, "Task #12 — qa-reviewer gains USER→403 audit log check  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #12.  Files touched: .claude/agents/qa-reviewer.md (one section).", italic=True)
add_subtask_table(doc, [
    ("12.1", "Add a new bullet under \"1. RBAC drift\" for the `GET /api/audit-logs` USER→403 check",
     "qa-reviewer.md (one bullet, one section)", "`grep -n \"USER.*403.*audit\" qa-reviewer.md` returns the new line", False),
    ("12.2 🔒", "Add a severity tag: the missing USER check is `blocker` if the endpoint leaks audit rows",
     "qa-reviewer.md (one inline edit)", "The new bullet's severity line says `blocker`", True),
    ("12.3", "Cross-reference the SPEC.md §4.5 wording in the check (so future drift between SPEC and reviewer is obvious)",
     "qa-reviewer.md (one line)", "The new bullet ends with `See SPEC.md §4.5.`", False),
])

# ---- Task #13 ----
add_heading(doc, "Task #13 — Rewrite rbac-redesign-spec.md as resolution doc  📄", level=2)
add_paragraph(doc, "Parent task: TASKS.md #13.  Files touched: rbac-redesign-spec.md (full rewrite).", italic=True)
add_subtask_table(doc, [
    ("13.1", "Write the new header: historical framing, link to v0.3 SPEC, link to TASKS.md",
     "rbac-redesign-spec.md (top 10 lines)", "Header includes the phrase \"Phase 2 design, transcribed\"", False),
    ("13.2", "Write §1 \"Why this exists\" — mark the original sketch as historical, not active",
     "rbac-redesign-spec.md (one section)", "§1 says \"treat this file as the resolution record, not a parallel spec\"", False),
    ("13.3", "Write §2 \"Decisions for the 6 open questions\" — one subsection per question, each citing the v0.3 SPEC section that answers it",
     "rbac-redesign-spec.md (one section, 6 subsections)", "All 6 sketch questions present, each with a v0.3 answer", False),
    ("13.4", "Write §3 \"USER audit visibility\" — explain why USER sees engine execution logs but never the audit log",
     "rbac-redesign-spec.md (one section)", "§3 cross-references SPEC.md §3.1 and §4.5", False),
])

# ---- Task #14a ----
add_heading(doc, "Task #14a — SPEC.md adds `mustChangePassword` + change-password endpoint  📄", level=2)
add_paragraph(doc, "Parent task: TASKS.md #14a.  Files touched: SPEC.md (multiple sections, one file).", italic=True)
add_subtask_table(doc, [
    ("14a.1", "§3.2 `User` entity: add `mustChangePassword: boolean` row to the field table",
     "SPEC.md (one row in §3.2 table)", "`grep -n \"mustChangePassword\" SPEC.md` returns the row in §3.2", False),
    ("14a.2", "§3.5 `AuditAction` enum: add `CHANGE_PASSWORD` to the enum block",
     "SPEC.md (one line in §3.5)", "`grep -n \"CHANGE_PASSWORD\" SPEC.md` returns the enum value", False),
    ("14a.3", "§4.2 `LoginResponse` gains `mustChangePassword: boolean`",
     "SPEC.md (one field in the LoginResponse shape)", "The LoginResponse JSON example includes `mustChangePassword`", False),
    ("14a.4 🔒", "§4.2 new endpoint: `POST /api/auth/change-password` with body, errors, and audit row documented",
     "SPEC.md (one new endpoint section)", "`grep -n \"POST /api/auth/change-password\" SPEC.md` returns the new section", True),
    ("14a.5", "§5.1 routes: new `/change-password` page",
     "SPEC.md (one row in the routes table)", "The routes table includes `/change-password`", False),
    ("14a.6", "§5.2 Login page: redirect to `/change-password` if `mustChangePassword = true`; new §5.2.1 Change Password page",
     "SPEC.md (one paragraph + one new page section)", "`grep -n \"Change Password\" SPEC.md` returns the new page section", False),
])

# ---- Task #14 ----
add_heading(doc, "Task #14 — Backend foundation: Gradle deps + Flyway V1__init.sql + 3 JPA entities  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #14.  Files touched: build.gradle, V1__init.sql, three entity files.", italic=True)
add_subtask_table(doc, [
    ("14.1", "Update `build.gradle` with the v0.3 dep set per SPEC §2.4",
     "build.gradle (one file)", "`./gradlew dependencies --configuration runtimeClasspath` lists `spring-boot-starter-data-jpa`, `jjwt-api`, `jasypt-spring-boot-starter`, `sshd-core`", False),
    ("14.2", "Create `V1__init.sql` with the `users` table (UUID PK, `mustChangePassword` deferred to #14b)",
     "db/migration/V1__init.sql (one table)", "`grep -c \"create table users\" V1__init.sql` returns `1`", False),
    ("14.3", "Create `V1__init.sql` `engines` table with all v0.3 columns including `serverPassword` and `mode`",
     "V1__init.sql (one table)", "`grep -c \"create table engines\" V1__init.sql` returns `1`", False),
    ("14.4", "Create `V1__init.sql` `audit_log` table with `jsonb` details + 3 indexes",
     "V1__init.sql (one table, one index block)", "`grep -c \"create table audit_log\" V1__init.sql` returns `1`, all 3 indexes present", False),
    ("14.5", "Create `V1__init.sql` `user_engine_access` join table",
     "V1__init.sql (one table)", "`grep -c \"create table user_engine_access\" V1__init.sql` returns `1`", False),
    ("14.6", "Create `User` JPA entity without `mustChangePassword` (added in #14b)",
     "…/user/User.java (one file)", "`./gradlew compileJava` green; `User` has UUID id, `@Version`, `passwordHash`, `roleType`, `assignedEngines` LAZY", False),
    ("14.7", "Create `Engine` JPA entity",
     "…/engine/EngineEntity.java (one file)", "`./gradlew compileJava` green; `Engine` has all v0.3 fields, no Lombok `@Data`", False),
    ("14.8", "Create `AuditLog` JPA entity (insert-only, no `@Version`)",
     "…/audit/AuditLog.java (one file)", "`./gradlew compileJava` green; `details` field is `@JdbcTypeCode(SqlTypes.JSON)`", False),
])

# ---- Task #14b ----
add_heading(doc, "Task #14b — `User` entity gains `mustChangePassword`  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #14b.  Files touched: User.java, V1__init.sql.  Depends on: #14.", italic=True)
add_subtask_table(doc, [
    ("14b.1 🔒", "Add `private boolean mustChangePassword;` with `@Column(nullable = false)` to `User`",
     "User.java (one field)", "`./gradlew compileJava` green; field exists with `@Column(nullable = false)`", True),
    ("14b.2", "Default `mustChangePassword = true` in the JPA constructor and `@PrePersist`",
     "User.java (one method)", "New `User` instances default to `mustChangePassword = true` (covered by the test in #15)", False),
    ("14b.3", "Add the `must_change_password BOOLEAN NOT NULL DEFAULT TRUE` column to `V1__init.sql`",
     "V1__init.sql (one column)", "The column appears in the `users` CREATE TABLE block", False),
])

# ---- Task #15 ----
add_heading(doc, "Task #15 — Backend: repositories + SecurityConfig + JWT filter + login  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #15.  Files touched: ~14 new files.  Depends on: #14, #14b.", italic=True)
add_subtask_table(doc, [
    ("15.1", "Create `UserRepository extends JpaRepository<User, UUID>` with `findByUsername(String)` and the role-checked query helpers",
     "…/user/UserRepository.java", "`./gradlew compileJava` green", False),
    ("15.2", "Create `EngineRepository extends JpaRepository<EngineEntity, UUID>` with `findByCodeAndDeletedAtIsNull(String)`",
     "…/engine/EngineRepository.java", "`./gradlew compileJava` green", False),
    ("15.3", "Create `AuditLogRepository extends JpaRepository<AuditLog, UUID>` with the filter-by-actor/engine/date query",
     "…/audit/AuditLogRepository.java", "`./gradlew compileJava` green", False),
    ("15.4 🔒", "Create `JasyptConfig` with `StringEncryptor` bean reading `JASYPT_ENCRYPTOR_PASSWORD` env var",
     "…/config/JasyptConfig.java", "App fails to start if the env var is missing", True),
    ("15.5", "Create `CorsConfig` (dev origin `http://localhost:5173` with `allowCredentials=true`) and `JacksonConfig` (`JavaTimeModule`, `NON_NULL`)",
     "…/config/CorsConfig.java, …/config/JacksonConfig.java", "App starts, CORS preflight from the dev origin returns the right headers", False),
    ("15.6 🔒", "Create `JwtService` (sign/validate, claims `sub`, `roles`, `mustChangePassword`)",
     "…/auth/JwtService.java", "`JwtService.sign(...)` returns a token that `JwtService.parse` can validate", True),
    ("15.7 🔒", "Create `UserPrincipal` (Spring `UserDetails` wrapping `User`)",
     "…/auth/UserPrincipal.java", "`UserPrincipal` exposes `getUsername()`, `getPassword()`, `getAuthorities()` from the entity", True),
    ("15.8 🔒", "Create `JwtAuthFilter` that extracts `Authorization: Bearer …` and parses it (parse step only — no security context population yet)",
     "…/auth/JwtAuthFilter.java", "A request with a valid token reaches the controller; one with a malformed token returns 401", True),
    ("15.9 🔒", "Extend `JwtAuthFilter` to populate `SecurityContext` with the `UserPrincipal` (authenticate step)",
     "…/auth/JwtAuthFilter.java (same file, second pass)", "A request with a valid token has `SecurityContextHolder` populated; `Authentication.getName()` returns the username", True),
    ("15.10", "Create `SecurityConfig` with the filter chain, JPA `UserDetailsService`, `@EnableMethodSecurity`, role hierarchy",
     "…/config/SecurityConfig.java", "`POST /api/auth/login` is permit-all; `GET /api/auth/me` requires authentication; `GET /api/engines` requires authentication", False),
    ("15.11", "Create `AuthController` with `POST /api/auth/login`, `POST /api/auth/logout`, `GET /api/auth/me`",
     "…/auth/AuthController.java", "`curl -X POST /api/auth/login -d '{...}'` returns a JWT and a `mustChangePassword` field", False),
    ("15.12 🔒", "Create `POST /api/auth/change-password` endpoint per SPEC §4.2 (the #14a endpoint, implemented here)",
     "…/auth/AuthController.java (one new method)", "`curl -X POST /api/auth/change-password` with the right JWT returns a new JWT with `mustChangePassword=false`", True),
    ("15.13 🔒", "Create `AuthenticationSuccessListener` and `AuthenticationFailureListener` for `LOGIN_SUCCESS` / `LOGIN_FAIL` audit rows",
     "…/audit/AuthenticationSuccessListener.java, …/audit/AuthenticationFailureListener.java", "A successful login writes a `LOGIN_SUCCESS` row; a failed login writes a `LOGIN_FAIL` row with `reason: \"BAD_CREDENTIALS\"`", True),
    ("15.14", "Add `V2__seed_admin.sql` (dev profile only) with one SYS_ADMIN (`mustChangePassword=true`, BCrypt-hashed password)",
     "db/seed/V2__seed_admin.sql (new file, dev profile only)", "`application-dev.properties` activates the seed location; running bootRun creates the admin row", False),
    ("15.15", "Write a `@SpringBootTest` smoke that boots the context, logs in as the seeded admin, and calls `GET /api/auth/me`",
     "…/src/test/…/AuthSmokeTest.java", "`./gradlew test` is green; the smoke passes", False),
])

# ---- Task #16 ----
add_heading(doc, "Task #16 — Backend: User CRUD endpoints + @Audited + AuditAspect  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #16.  Files touched: ~8 new files.  Depends on: #15.", italic=True)
add_subtask_table(doc, [
    ("16.1", "Create `AuditAction` enum with the full set (12 values)",
     "…/audit/AuditAction.java", "`./gradlew compileJava` green; all 12 values present", False),
    ("16.2", "Create `@Audited` annotation (`action`, `targetEngineFromPath`, `details` SpEL)",
     "…/audit/Audited.java", "The annotation compiles and is usable on a method", False),
    ("16.3 🔒", "Create `AuditAspect` (AOP): extract actor from `SecurityContext`, resolve target engine from path, write row on success AND on exception",
     "…/audit/AuditAspect.java", "A `@Audited` controller method produces exactly one `AuditLog` row per call (success or failure)", True),
    ("16.4", "Create DTOs: `CreateUserRequest`, `UpdateUserRolesRequest`, `UserResponse` with validation annotations",
     "…/user/dto/*.java (three files)", "Validation rejects missing/blank fields; `UserResponse` does not include `passwordHash`", False),
    ("16.5", "Create `UserService` (BCrypt hash, set `mustChangePassword = true` on create, prevent last-SYS_ADMIN delete, prevent self-delete)",
     "…/user/UserService.java", "Integration test verifies self-delete, last-SYS_ADMIN, password hash, and short-password rules", False),
    ("16.6", "Create `UserController` `GET /api/users` (SYS_ADMIN + ADMIN)",
     "…/user/UserController.java (one method)", "`curl` as SYS_ADMIN returns all users; as ADMIN returns all users; as USER returns 403", False),
    ("16.7 🔒", "Create `UserController` `POST /api/users` with the role-checked rule (ADMIN can only create USER; SYS_ADMIN can create any)",
     "…/user/UserController.java (one method)", "`curl` as ADMIN with `roleType=ADMIN` returns 403; as ADMIN with `roleType=USER` returns 201", True),
    ("16.8", "Create `UserController` `DELETE /api/users/{id}` with the self-delete and last-SYS_ADMIN guards",
     "…/user/UserController.java (one method)", "`curl` as a SYS_ADMIN trying to delete themselves returns 400; trying to delete the last SYS_ADMIN returns 400", False),
    ("16.9 🔒", "Create `UserController` `PATCH /api/users/{id}/roles` with the role-checked rule (ADMIN can only change USER; SYS_ADMIN any)",
     "…/user/UserController.java (one method)", "`curl` as ADMIN trying to promote a USER to ADMIN returns 403", True),
    ("16.10", "Write `@SpringBootTest` cases for each endpoint covering success + each failure mode",
     "…/src/test/…/UserControllerTest.java", "`./gradlew test` is green; the 4 endpoint cases pass", False),
])

# ---- Task #17 ----
add_heading(doc, "Task #17 — Backend: Engine CRUD + factory + MOCK impl  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #17.  Files touched: ~9 new files.  Depends on: #15.", italic=True)
add_subtask_table(doc, [
    ("17.1", "Create the `OrderEngineOperations` interface (preserved from v0.2)",
     "…/engine/OrderEngineOperations.java", "The interface compiles; the 7 methods match SPEC §3.6", False),
    ("17.2", "Create the `EngineStatus` and `EngineMode` enums",
     "…/engine/EngineStatus.java, …/engine/EngineMode.java", "The enums compile and are referenced by the entity (already in #14)", False),
    ("17.3", "Create `EngineNotSupportedException` (→ 404), `EngineUnreachableException` (→ 502), `EngineAuthException` (→ 403), `EngineScriptException` (→ 502, carries `exitCode` and `stderr`)",
     "…/engine/*.java (four files)", "All four exceptions compile; mapped to HTTP codes by `ApiExceptionHandler` (created in #18)", False),
    ("17.4", "Create the `LogLine` value type (timestamp, level, message)",
     "…/engine/LogLine.java", "The type compiles", False),
    ("17.5", "Create `MockEngineOperations` (in-memory state machine; per-engine wrapper via the factory)",
     "…/engine/impl/MockEngineOperations.java", "A `MOCK` engine's `start()` transitions STOPPED → RUNNING; `getLogs(100)` returns the last 100 lines", False),
    ("17.6", "Create `OrderEngineFactory` (looks up by `code` from `EngineRepository.findByCodeAndDeletedAtIsNull`, throws `EngineNotSupportedException` on miss)",
     "…/engine/OrderEngineFactory.java", "Unknown code throws → 404; `MOCK` returns wrapper; `REAL` returns `SshBackedEngine` (created in #19)", False),
    ("17.7", "Create `EngineService` (CRUD: create, soft-delete with cascade, PATCH ssh; emits `EngineStatusChangedEvent`)",
     "…/engine/EngineService.java", "Created engine appears in DB; soft-deleted engine is no longer returned; event fires on status transitions", False),
    ("17.8", "Create `EngineController` with the 4 CRUD endpoints (`GET /api/engines`, `POST /api/engines`, `DELETE /api/engines/{code}`, `PATCH /api/engines/{code}/ssh`)",
     "…/engine/EngineController.java", "`curl` as SYS_ADMIN can create, soft-delete (204), PATCH ssh; as USER returns 403 on the create/delete/patch", False),
    ("17.9", "Write `@SpringBootTest` cases for the 4 CRUD endpoints",
     "…/src/test/…/EngineControllerTest.java", "`./gradlew test` green", False),
])

# ---- Task #18 ----
add_heading(doc, "Task #18 — Backend: engine status/start/stop/logs endpoints  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #18.  Files touched: EngineController, ApiExceptionHandler, EngineService.  Depends on: #16, #17.", italic=True)
add_subtask_table(doc, [
    ("18.1", "Create `ApiExceptionHandler` (`@ControllerAdvice`) mapping the 4 engine exceptions to their HTTP codes per `ssh-engine-ops.md`",
     "…/web/ApiExceptionHandler.java", "A thrown `EngineAuthException` produces 403, `EngineUnreachableException` produces 502, `EngineScriptException` produces 502, future-timeout produces 504", False),
    ("18.2 🔒", "Add `GET /api/engines/{code}/status` to `EngineController` with `@PreAuthorize` and the assignedEngines filter for USER",
     "EngineController.java (one method)", "`curl` as USER for an assigned engine returns 200; for an unassigned engine returns 403", True),
    ("18.3 🔒", "Add `POST /api/engines/{code}/start` to `EngineController` with `@Audited(action = START_ENGINE, targetEngineFromPath = true)` and `@PreAuthorize`",
     "EngineController.java (one method)", "Successful start writes a `START_ENGINE` row with `exitCode: 0`; failed start writes a row with the error in `details`", True),
    ("18.4 🔒", "Add `POST /api/engines/{code}/stop` to `EngineController`, mirror of #18.3",
     "EngineController.java (one method)", "Same as #18.3 for `STOP_ENGINE`", True),
    ("18.5", "Add `GET /api/engines/{code}/logs?limit=N` to `EngineController` with validation on the `limit` enum (`{50, 100, 200}`)",
     "EngineController.java (one method)", "`curl ?limit=42` returns 400; `?limit=100` returns the last 100 lines from `LogBuffer`", False),
    ("18.6", "Write `@SpringBootTest` cases for each of the 4 endpoints with both happy path and the exception-mapping cases",
     "…/src/test/…/EngineActionsTest.java", "`./gradlew test` green", False),
    ("18.7", "Add `EngineStatusChangedEvent` + Spring `ApplicationEventPublisher` wiring in `EngineService` so the WS handler in #20 can subscribe",
     "…/engine/EngineService.java (one method)", "A successful start publishes the event; a test listener receives it", False),
])

# ---- Task #19 ----
add_heading(doc, "Task #19 — Backend: SshBackedEngine + background log tailer  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #19.  Files touched: SshBackedEngine, SshClientProvider, LogBuffer, LogTailerRegistry.  Depends on: #17, #18.", italic=True)
add_subtask_table(doc, [
    ("19.1 🔒", "Create `SshClientProvider`: per-engine cached `SshClient`, idle-evicted at 5 min via a `ScheduledExecutorService`, closed in `@PreDestroy`",
     "SshClientProvider.java", "Two calls to the same engine within 5 min reuse one `SshClient`; a third call after 6 min opens a new one", True),
    ("19.2 🔒", "Create `SshBackedEngine.status()` with the 5s connect + 5s op timeouts; throws `EngineUnreachableException` on connect fail (with one retry)",
     "SshBackedEngine.java (one method)", "A `status()` against a reachable host returns `RUNNING`/`STOPPED`; against an unreachable host throws within 5s, then retries once", True),
    ("19.3a 🔒", "Create `SshBackedEngine.start()` with the 30s bound via `Future.get(30, SECONDS)`; throws `EngineScriptException(exitCode, stderr)` on non-zero exit",
     "SshBackedEngine.java (one method)", "A 60s-hanging start script is cancelled at 30s and returns 504; a start script that exits 127 returns 502 with the stderr in `details`", True),
    ("19.3b 🔒", "Create `SshBackedEngine.stop()` with the same 30s bound; mirror of start for the running → stopped transition",
     "SshBackedEngine.java (one method)", "A 60s-hanging stop script is cancelled at 30s and returns 504; a stop script that exits 127 returns 502 with the stderr in `details`", True),
    ("19.4", "Create `SshBackedEngine.getLogs(int)` with the 10s bound; returns the last N lines via a single `tail -n N` invocation",
     "SshBackedEngine.java (one method)", "A request for 100 lines returns 100 lines within 10s on a normal host", False),
    ("19.5 🔒", "Create `LogBuffer`: per-engine `ArrayDeque<LogLine>`, cap 500, oldest evicted on insert",
     "LogBuffer.java", "Inserting a 501st line evicts the oldest; the buffer is per-engine (two engines don't share)", True),
    ("19.6", "Create the SSH channel reader loop (one method) that reads a `ChannelExec` stdout line-by-line and pushes to the buffer",
     "SshBackedEngine.java (one method)", "A `tail -F` channel's lines appear in the buffer in real time", False),
    ("19.7 🔒", "Create `LogTailerRegistry` listening to `EngineStatusChangedEvent`: starts a tailer thread on `RUNNING` for `mode=REAL` engines, stops on `STOPPED`/deletion/5 reconnect failures",
     "LogTailerRegistry.java", "A `RUNNING` REAL engine has a tailer thread; a `STOPPED` engine does not; 5 reconnect failures emits a single `WARN` and stops", True),
    ("19.8", "Write integration tests against a `sshd` Maven-embedded server (covers the 4 SSH operations end to end)",
     "…/src/test/…/SshBackedEngineIT.java", "`./gradlew test` green; the 4 ops are exercised against a real `sshd` instance", False),
])

# ---- Task #20 ----
add_heading(doc, "Task #20 — Backend: WebSocket logs/stream handler  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #20.  Files touched: EngineLogsWebSocketHandler, WebSocketSessionRegistry, WebSocketConfig, SecurityConfig.  Depends on: #15, #18, #19.", italic=True)
add_subtask_table(doc, [
    ("20.1", "Add `spring-boot-starter-websocket` to `build.gradle` (already on the classpath from #14) and the `WebSocketConfigurer` bean",
     "WebSocketConfig.java (new)", "A client can open a WebSocket to the configured path", False),
    ("20.2", "Create `WebSocketSessionRegistry`: per-engine set of sessions, `add(code, session)`, `remove(code, session)`, `broadcast(code, line)`",
     "WebSocketSessionRegistry.java", "A line broadcast to engine `bpl` reaches all sessions registered for `bpl` and no others", False),
    ("20.3 🔒", "Create `EngineLogsWebSocketHandler.afterConnectionEstablished`: snapshot last 100 lines from `LogBuffer`, send them as JSON, register the session",
     "EngineLogsWebSocketHandler.java (one method)", "A new connection receives the snapshot before any live line", True),
    ("20.4", "Create `EngineLogsWebSocketHandler.handleMessage` (client pings; currently no-op beyond keep-alive)",
     "EngineLogsWebSocketHandler.java (one method)", "Client pings don't crash; the server sends no data in response (snapshot is on connect)", False),
    ("20.5", "Create `EngineLogsWebSocketHandler.afterConnectionClosed`: remove the session from the registry",
     "EngineLogsWebSocketHandler.java (one method)", "A closed client is no longer in the registry; subsequent broadcasts don't reach it", False),
    ("20.6 🔒", "JWT auth on the WebSocket handshake via the existing `JwtAuthFilter`: extend `SecurityConfig` to require authentication for the WS path and reject the handshake with 401 if the token is missing/invalid",
     "SecurityConfig.java (extend), WebSocketConfig.java (wire)", "A WS handshake without a token returns 401; one with a valid token proceeds", True),
    ("20.7", "Role + assignment gate on the WS path: reject with 403 if the caller is a USER without the engine in their `assignedEngines`",
     "EngineLogsWebSocketHandler.java (one method)", "A USER's WS connection to a non-assigned engine is rejected with 403; a USER's connection to an assigned engine succeeds", False),
    ("20.8", "Wire the `LogTailerRegistry` (from #19) to also push into the WS session registry: every tailer line is broadcast to live sessions AND pushed to the buffer",
     "LogTailerRegistry.java (extend)", "A live `tail -F` line appears in both the buffer and every connected viewer within 100ms", False),
])

# ---- Task #21 ----
add_heading(doc, "Task #21 — Backend: /api/audit-logs endpoint  ⚙️", level=2)
add_paragraph(doc, "Parent task: TASKS.md #21.  Files touched: AuditLogController, AuditService, AuditLogResponse.  Depends on: #16.", italic=True)
add_subtask_table(doc, [
    ("21.1", "Create `AuditLogResponse` DTO with `details` as a parsed `Map<String, Object>` (not a string)",
     "…/audit/dto/AuditLogResponse.java", "The DTO compiles; `details` round-trips as a JSON object", False),
    ("21.2 🔒", "Create `AuditService` query method: filters by `actor`, `action`, `engine` (resolves code to UUID), `from`/`to` (Instant), with paginated `Pageable`",
     "AuditService.java", "A query with all 5 filters returns the expected subset; an out-of-range `from` after `to` returns 400", True),
    ("21.3 🔒", "Create `AuditLogController.GET /api/audit-logs` with `@PreAuthorize(\"hasAnyRole('SYS_ADMIN', 'ADMIN')\")` — USER is rejected outright, not filtered",
     "AuditLogController.java (one method)", "`curl` as USER returns 403; as ADMIN returns the full system log", True),
    ("21.4", "Write `@SpringBootTest` cases for the 3 filter dimensions and the USER→403 case",
     "…/src/test/…/AuditLogControllerTest.java", "`./gradlew test` green", False),
])

# ---- Task #22 ----
add_heading(doc, "Task #22 — Backend: build + test pass + DB up  ✅", level=2)
add_paragraph(doc, "Parent task: TASKS.md #22.  Files touched: docker-compose.yml (new), no app code. Verification only.  Depends on: #14–#21.", italic=True)
add_subtask_table(doc, [
    ("22.1", "Create `docker-compose.yml` for local Postgres (`postgres:16-alpine`, port 5432, volume for data)",
     "docker-compose.yml (new)", "`docker compose up -d` brings up Postgres; `psql` connects", False),
    ("22.2", "Set the env vars in a `.env.local` (gitignored) for the local run: `DB_URL`, `DB_USERNAME`, `DB_PASSWORD`, `JWT_SECRET`, `JASYPT_ENCRYPTOR_PASSWORD`",
     ".env.local (new, gitignored)", "`source .env.local && ./gradlew bootRun` starts the app without missing-env errors", False),
    ("22.3", "`./gradlew compileJava` clean", "(no file)", "Exit 0", False),
    ("22.4", "`./gradlew test` clean (all backend test classes from #15–#21)", "(no file)", "All tests pass", False),
    ("22.5", "`./gradlew bootRun` starts; verify `GET /actuator/health` is `UP`", "(no file)", "`curl /actuator/health` returns `{\"status\":\"UP\"}`", False),
    ("22.6", "Curl as SYS_ADMIN: every endpoint returns the expected 2xx; verify `@Audited` writes one `AuditLog` row per state change", "(no file)", "The `audit_log` table has the expected rows after the curl session", False),
    ("22.7", "Curl as ADMIN: full audit access; can create only USER-role users; cannot delete another ADMIN", "(no file)", "Admin endpoints return the expected 2xx/403 mix", False),
    ("22.8", "Curl as USER: 403 on `/api/audit-logs` (per #11); empty engine list if no assignments; only assigned engines visible", "(no file)", "USER requests return the expected 403/200/[] mix", False),
    ("22.9", "WebSocket handshake with a USER JWT for an assigned engine returns 200 + initial snapshot; for an unassigned engine returns 403", "(no file)", "`wscat` against assigned code accepts and delivers snapshot; against unassigned returns 403", False),
    ("22.10", "`qa-reviewer` pass on the full backend diff (#14–#21)", "(no file)", "qa-reviewer reports no `blocker` findings", False),
])

# ---- Task #23 ----
add_heading(doc, "Task #23 — Frontend: AuthContext + api/client + router shell  🎨", level=2)
add_paragraph(doc, "Parent task: TASKS.md #23.  Files touched: AuthContext.tsx, api/client.ts, App.tsx, AppShell.tsx, main.tsx.  Depends on: #15.", italic=True)
add_subtask_table(doc, [
    ("23.1 🔒", "Create `api/client.ts`: `fetch` wrapper with `Authorization: Bearer …`, 401 → clear token + redirect to `/login`, error envelope unwrap to thrown `Error(message)`",
     "src/api/client.ts", "A request with a valid token succeeds; one with an expired token redirects to `/login`; the thrown `Error` message matches the server's `message` field", True),
    ("23.2 🔒", "Create `AuthContext`: state `{ user, token, mustChangePassword, isLoading }`; on mount, validate token via `GET /api/auth/me`; on 401, clear + redirect",
     "src/auth/AuthContext.tsx", "The context provides the user; an expired token clears the user and redirects; a fresh token hydrates the user", True),
    ("23.3", "Wire `AuthContext.login(username, password)`: POST `/api/auth/login`, store token, set user",
     "src/auth/AuthContext.tsx (one method)", "A successful login sets `user` and `token`; an invalid login throws an `Error(\"Invalid credentials\")`", False),
    ("23.4", "Wire `AuthContext.logout()`: clear token, clear user, redirect to `/login`",
     "src/auth/AuthContext.tsx (one method)", "A click on the logout button redirects to `/login` and the next protected-route render bounces back to `/login`", False),
    ("23.5", "Create `App.tsx` with the React Router routes: `/login`, `/dashboard`, `/logs`, `/admin` (lazy for non-USER), `/change-password`, `/404`, `/403`",
     "src/App.tsx", "All 7 routes resolve; `/admin` is not in the bundle when the user role is USER (verified by inspecting the build output)", False),
    ("23.6", "Create `AppShell` (top bar with role badge + logout button) used by all authenticated pages",
     "src/components/AppShell.tsx", "A logged-in SYS_ADMIN sees the admin link; a USER does not; both see their role badge", False),
    ("23.7", "`npm run build` clean; `npm run lint` clean", "(no file)", "Build passes; no lint errors", False),
])

# ---- Task #24 ----
add_heading(doc, "Task #24 — Frontend: Login page  🎨", level=2)
add_paragraph(doc, "Parent task: TASKS.md #24.  Files touched: src/pages/Login.tsx, src/auth/AuthContext.tsx.  Depends on: #23.", italic=True)
add_subtask_table(doc, [
    ("24.1", "Create `Login.tsx`: username + password form, submit handler, inline error display",
     "src/pages/Login.tsx", "The form renders; submit triggers `AuthContext.login`; an error renders inline", False),
    ("24.2", "Wire the post-login redirect: `/change-password` if `mustChangePassword = true`, else `/dashboard`",
     "Login.tsx (or `AuthContext.login`)", "A login where the server returns `mustChangePassword: true` navigates to `/change-password`; one with `false` goes to `/dashboard`", False),
    ("24.3 🔒", "Inline \"Invalid credentials\" on 401, with no field-level enumeration",
     "Login.tsx", "A bad password shows a single message; a missing user shows the same message", True),
    ("24.4", "`npm run build` clean", "(no file)", "Build passes", False),
])

# ---- Task #24a ----
add_heading(doc, "Task #24a — Frontend: ChangePassword page  🎨", level=2)
add_paragraph(doc, "Parent task: TASKS.md #24a.  Files touched: src/pages/ChangePassword.tsx, src/App.tsx, src/auth/AuthContext.tsx.  Depends on: #15, #23.", italic=True)
add_subtask_table(doc, [
    ("24a.1 🔒", "Create `ChangePassword.tsx` form: current password, new password, confirm new password, with client-side \"new matches confirm\" check",
     "src/pages/ChangePassword.tsx", "The form renders; mismatched confirm shows an inline error before submit", True),
    ("24a.2 🔒", "Wire submit to `POST /api/auth/change-password`; on success, store the new JWT, clear `mustChangePassword`, redirect to `/dashboard`",
     "ChangePassword.tsx (one method)", "Successful change navigates to `/dashboard`; a 401 (bad current) shows the right error; a 422 shows the validation message", True),
    ("24a.3", "Register the `/change-password` route in `App.tsx` (authenticated-only)",
     "src/App.tsx (one line)", "An unauthenticated visit to `/change-password` redirects to `/login`; an authenticated visit with `mustChangePassword = true` renders the form", False),
    ("24a.4", "`npm run build` clean", "(no file)", "Build passes", False),
])

# ---- Task #25 ----
add_heading(doc, "Task #25 — Frontend: Dashboard with EngineCards + WS + polling  🎨", level=2)
add_paragraph(doc, "Parent task: TASKS.md #25.  Files touched: Dashboard.tsx, EngineCard.tsx, StatusPill.tsx, useEngineLogsSocket.ts, useEngineStatus.ts.  Depends on: #23, #18, #19, #20.", italic=True)
add_subtask_table(doc, [
    ("25.1", "Create `useEngineLogsSocket` hook: WS connect with JWT, exponential backoff reconnect (1s → 30s cap), clean teardown on unmount",
     "src/hooks/useEngineLogsSocket.ts", "A closed network reconnects with backoff; unmounting the component closes the socket", False),
    ("25.2", "Create `useEngineStatus` hook: WS primary, polling fallback (every 5s) when WS is closed",
     "src/hooks/useEngineStatus.ts", "With WS open, the hook reports the live status; with WS closed, it polls and reports every 5s", False),
    ("25.3", "Create `StatusPill` component (color-coded pill: green=RUNNING, gray=STOPPED, red=ERROR)",
     "src/components/StatusPill.tsx", "The pill renders the right color for each status", False),
    ("25.4", "Create `EngineCard` component: name, code, `StatusPill`, `lastTransitionAt`, Start/Stop buttons (role-gated, in-flight disabled)",
     "src/components/EngineCard.tsx", "A USER sees Start/Stop only for assigned engines; ADMIN/SYS_ADMIN see them for all visible; in-flight buttons are disabled", False),
    ("25.5", "Create `Dashboard.tsx`: grid of `EngineCard` per visible engine, filtered by `currentUser.assignedEngines`; SYS_ADMIN sees \"+ Add Engine\"",
     "src/pages/Dashboard.tsx", "The dashboard shows the right cards for each role; the \"+ Add Engine\" button is SYS_ADMIN-only", False),
    ("25.6", "Wire `EngineCard` to `useEngineStatus` and to the start/stop `POST` actions",
     "src/components/EngineCard.tsx (extend)", "A click on Start transitions the card to RUNNING within 5s; a failure surfaces the error message", False),
    ("25.7 🔒", "Verify the dashboard never renders a card for an engine the user is not assigned to (defense in depth — the server filters, but the UI must not flash unassigned engines during the role transition)",
     "src/pages/Dashboard.tsx (one filter line)", "A USER with no assignments sees an empty state, not a 403 error; a USER who loses assignment mid-session doesn't see the engine card on next render", True),
    ("25.8", "`npm run build` clean; `npm run lint` clean", "(no file)", "Build passes", False),
])

# ---- Task #26 ----
add_heading(doc, "Task #26 — Frontend: Logs page with both filter types  🎨", level=2)
add_paragraph(doc, "Parent task: TASKS.md #26.  Files touched: Logs.tsx, AuditLogsTable.tsx, EngineLogsTable.tsx.  Depends on: #21, #23, #25.", italic=True)
add_subtask_table(doc, [
    ("26.1", "Create `AuditLogsTable` component: paginated table with `actor`, `action`, `targetEngineCode`, `timestamp`, `details` (raw-JSON toggle)",
     "src/components/AuditLogsTable.tsx", "A page of audit logs renders; the raw-JSON toggle shows the parsed `details` object", False),
    ("26.2", "Create `EngineLogsTable` component: virtualized list of `LogLine`s, fed by both the initial `GET /api/engines/{code}/logs?limit=100` snapshot and the WS stream",
     "src/components/EngineLogsTable.tsx", "The table shows the initial 100 lines, then live updates as the WS pushes; no duplicates from the snapshot + first live lines", False),
    ("26.3 🔒", "Filter the Source dropdown: USER sees only \"Engine Execution Logs\"; ADMIN/SYS_ADMIN see both options",
     "src/pages/Logs.tsx", "A USER visiting `/logs` cannot select \"System Audit Logs\" — the option is not in the dropdown", True),
    ("26.4", "Filter dropdown for Engine: shows all visible engines; selecting one switches the table to engine execution logs",
     "src/pages/Logs.tsx", "Selecting engine `bpl` switches the table to its `LogBuffer` + WS", False),
    ("26.5", "Query params for audit logs (`actor`, `action`, `engine`, `from`, `to`, `page`, `size`) wired to filter controls",
     "src/pages/Logs.tsx", "Setting a filter updates the URL query and the table", False),
    ("26.6", "Wire the WS subscription for the selected engine: when the user changes engine, close the old socket and open a new one",
     "src/pages/Logs.tsx (one effect)", "Switching engines closes the old socket and opens a new one; the table switches without orphaning the old subscription", False),
    ("26.7", "`npm run build` clean; `npm run lint` clean", "(no file)", "Build passes", False),
])

# ---- Task #27 ----
add_heading(doc, "Task #27 — Frontend: Admin Panel (Users + Engines tabs)  🎨", level=2)
add_paragraph(doc, "Parent task: TASKS.md #27.  Files touched: Admin.tsx, UserForm.tsx, EngineForm.tsx.  Depends on: #16, #17, #23.", italic=True)
add_subtask_table(doc, [
    ("27.1", "Create `Admin.tsx` with two tabs (Users, Engines); route is lazy-imported for non-USER",
     "src/pages/Admin.tsx", "The page renders; the Users tab is the default; the Engines tab is SYS_ADMIN-only", False),
    ("27.2", "Create the Users tab: table with `username`, `role`, `assignedEngineCodes`, [Edit] [Delete]; the table is the same for ADMIN and SYS_ADMIN, with row-action restrictions",
     "src/pages/Admin.tsx (one component)", "A USER row's [Edit]/[Delete] is enabled for both ADMIN and SYS_ADMIN; an ADMIN/SYS_ADMIN row's [Edit]/[Delete] is enabled only for SYS_ADMIN", False),
    ("27.3", "Create `UserForm.tsx`: username, password, role (constrained by caller's role per SPEC §3.1), assignedEngines multi-select; used for both Add and Edit",
     "src/components/UserForm.tsx", "An ADMIN's role select does not include `ADMIN` or `SYS_ADMIN`; a SYS_ADMIN's does", False),
    ("27.4", "Create the Engines tab (SYS_ADMIN only): table with `name`, `code`, `mode`, `serverIp`, [Edit SSH] [Delete]",
     "src/pages/Admin.tsx (one component)", "ADMIN sees the tab as read-only; SYS_ADMIN sees the full actions", False),
    ("27.5", "Create `EngineForm.tsx`: name, code, mode, serverIp, serverUsername, serverPassword (write-only), startScript, stopScript, logScript; used for both Add and Edit",
     "src/components/EngineForm.tsx", "The form sends a `POST` on Add and a `PATCH` on Edit; the password field is a write-only text input", False),
    ("27.6", "Wire the [Delete] actions to soft-delete with a confirm dialog (engines) or a hard delete with a confirm dialog (users)",
     "src/pages/Admin.tsx", "A click on [Delete] opens a confirm dialog; on confirm, the row disappears and a refresh shows the new state", False),
    ("27.7 🔒", "Ensure the serverPassword is **never** returned in the response that populates the form (no `defaultValue` from a fetched row, since the field doesn't exist on the DTO)",
     "src/components/EngineForm.tsx", "Opening [Edit SSH] on an existing engine shows the existing IP/username/scripts but an empty password field — the user must retype it to change it", True),
    ("27.8", "Wire the \"+ Add Engine\" button on the Dashboard to open the same `EngineForm` for creation",
     "src/pages/Dashboard.tsx (extend)", "A click on \"+ Add Engine\" opens the form modal; submit creates the engine and refreshes the dashboard", False),
    ("27.9", "`npm run build` clean; `npm run lint` clean", "(no file)", "Build passes", False),
])

# ---- Task #28 ----
add_heading(doc, "Task #28 — Frontend: build + manual role smoke + screenshots  ✅", level=2)
add_paragraph(doc, "Parent task: TASKS.md #28.  Files touched: no new files; verification only.  Depends on: #23–#27.", italic=True)
add_subtask_table(doc, [
    ("28.1", "`npm run build` clean", "(no file)", "Build passes", False),
    ("28.2", "`npm run lint` clean", "(no file)", "Lint passes", False),
    ("28.3", "Manual smoke as SYS_ADMIN: every page (Dashboard, Logs, Admin → Users, Admin → Engines) renders and the expected actions work", "(no file)", "All 4 pages render; adding an engine, adding a user, assigning a user all work end-to-end", False),
    ("28.4", "Manual smoke as ADMIN: Admin Panel visible, Engines tab is read-only, Users tab works for USER-role rows only", "(no file)", "ADMIN cannot create ADMIN/SYS_ADMIN; the engine SSH fields are not editable", False),
    ("28.5", "Manual smoke as USER: Admin Panel is not in the bundle; Logs page Source dropdown hides \"System Audit Logs\"; only assigned engines appear on the dashboard", "(no file)", "The build output for USER's bundle does not include the admin chunk; the Logs Source dropdown has one option", False),
    ("28.6", "Screenshots: capture Login, Dashboard (SYS_ADMIN), Dashboard (USER with assigned engines), Logs (audit view), Logs (engine view), Admin → Users, Admin → Engines (Add modal)",
     "docs/screenshots/*.png (7 files)", "7 PNGs in `docs/screenshots/` matching the slide deck layout", False),
    ("28.7", "`qa-reviewer` pass on the full frontend diff (#23–#27)", "(no file)", "qa-reviewer reports no `blocker` findings", False),
])

# ---- Summary ----
add_heading(doc, "Summary", level=1)
add_paragraph(doc, "Per-task subtask counts and security-seam (🔒) counts:")
add_summary_table(doc, [
    ("#11", 5, 0),
    ("#12", 3, 1),
    ("#13", 4, 0),
    ("#14a", 6, 1),
    ("#14", 8, 0),
    ("#14b", 3, 1),
    ("#15", 15, 7),
    ("#16", 10, 3),
    ("#17", 9, 0),
    ("#18", 7, 3),
    ("#19", 9, 6),
    ("#20", 8, 2),
    ("#21", 4, 2),
    ("#22", 10, 0),
    ("#23", 7, 2),
    ("#24", 4, 1),
    ("#24a", 4, 2),
    ("#25", 8, 1),
    ("#26", 7, 1),
    ("#27", 9, 1),
    ("#28", 7, 0),
    ("Total", 147, 34),
])

add_paragraph(
    doc,
    "Every TASKS.md task has a subtask table in TASKS-decomposed.md. Each subtask has a "
    "`done when` that's paste-able into a terminal or test file. When a subagent finishes "
    "a task, qa-reviewer cross-checks the implementation against the planned subtasks.",
    italic=True,
)

# ---- Save ----
output_path = "d:/BPL-Order-Engine-Admin/SUBTASKS.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
